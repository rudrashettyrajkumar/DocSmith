"""Builds the final polished Word (.docx) document with python-docx."""

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "outputs"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)  # corporate dark blue
GREY = RGBColor(0x59, 0x59, 0x59)


def _slugify(text: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
    return slug[:60] or "document"


def _style_headings(doc: Document) -> None:
    for level, size in (("Heading 1", 15), ("Heading 2", 12)):
        style = doc.styles[level]
        style.font.color.rgb = ACCENT
        style.font.size = Pt(size)


def build_docx(ctx: dict) -> str:
    """Render the agent's context (title, assumptions, sections) into a .docx."""
    OUTPUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    _style_headings(doc)

    # --- cover block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(ctx["title"])
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"{ctx['document_type']}  |  Audience: {ctx['audience']}  |  {date.today():%d %B %Y}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = GREY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Prepared automatically by the Autonomous Document Agent")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = GREY

    # --- assumptions the agent made (transparency for the reader)
    if ctx.get("assumptions"):
        doc.add_heading("Assumptions", level=2)
        for assumption in ctx["assumptions"]:
            doc.add_paragraph(assumption, style="List Bullet")

    # --- body sections
    for section in ctx.get("sections", []):
        doc.add_heading(section["heading"], level=1)
        for paragraph in section.get("paragraphs", []):
            doc.add_paragraph(paragraph)
        for bullet in section.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")

    # --- supporting data table
    if ctx.get("facts"):
        doc.add_heading("Appendix: Key Data Points", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        header = table.rows[0].cells
        header[0].text, header[1].text = "#", "Data point"
        for i, fact in enumerate(ctx["facts"], start=1):
            row = table.add_row().cells
            row[0].text, row[1].text = str(i), fact

    filename = f"{_slugify(ctx['title'])}-{date.today():%Y%m%d}.docx"
    doc.save(OUTPUT_DIR / filename)
    return filename
